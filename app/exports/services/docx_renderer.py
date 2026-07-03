from __future__ import annotations

from io import BytesIO

from app.compliance.renderers.docx_report_renderer import DocxReportRenderer
from app.exports.services.export_content_builder import ExportDocumentContent


class DocxRenderer:
    def render(self, content: ExportDocumentContent) -> bytes:
        try:
            from docx import Document  # type: ignore
        except Exception:  # noqa: BLE001
            return DocxReportRenderer._fallback_docx_bytes(content.title)

        doc = Document()
        doc.add_heading(content.branding.company_display_name, 0)
        doc.add_heading(content.title, 1)
        if content.subtitle:
            doc.add_paragraph(content.subtitle)
        doc.add_paragraph(f"Generated at: {content.generated_at.isoformat()}")

        for section in content.sections:
            doc.add_heading(section.title, 2)
            for key, value in section.rows:
                doc.add_paragraph(f"{key}: {value}")
            for paragraph in section.paragraphs:
                doc.add_paragraph(paragraph)
            for item in section.items:
                doc.add_paragraph(item, style="List Bullet")

        doc.add_paragraph(content.branding.footer_text)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
