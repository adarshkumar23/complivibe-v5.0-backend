import uuid
from datetime import UTC, datetime
from io import BytesIO
from typing import Any
from zipfile import ZIP_DEFLATED, ZipFile

from fastapi import HTTPException, status
from sqlalchemy.orm import Session

from app.compliance.renderers.report_section_mapper import ReportSectionMapper
from app.models.compliance_report import ComplianceReport
from app.models.compliance_report_section import ComplianceReportSection
from app.models.organization import Organization
from app.repositories.report_repository import ReportRepository
from app.services.report_service import REPORT_CAVEAT


class DocxReportRenderer:
    @staticmethod
    def _fallback_docx_bytes(text: str) -> bytes:
        # Minimal valid docx (ZIP package) fallback when python-docx is unavailable at runtime.
        content_types = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<Types xmlns=\"http://schemas.openxmlformats.org/package/2006/content-types\">"
            "<Default Extension=\"rels\" ContentType=\"application/vnd.openxmlformats-package.relationships+xml\"/>"
            "<Default Extension=\"xml\" ContentType=\"application/xml\"/>"
            "<Override PartName=\"/word/document.xml\" ContentType=\"application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml\"/>"
            "</Types>"
        )
        rels = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<Relationships xmlns=\"http://schemas.openxmlformats.org/package/2006/relationships\">"
            "<Relationship Id=\"rId1\" Type=\"http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument\" Target=\"word/document.xml\"/>"
            "</Relationships>"
        )
        doc_xml = (
            "<?xml version=\"1.0\" encoding=\"UTF-8\" standalone=\"yes\"?>"
            "<w:document xmlns:w=\"http://schemas.openxmlformats.org/wordprocessingml/2006/main\">"
            "<w:body><w:p><w:r><w:t>"
            + text
            + "</w:t></w:r></w:p><w:sectPr/></w:body></w:document>"
        )

        buffer = BytesIO()
        with ZipFile(buffer, "w", ZIP_DEFLATED) as zf:
            zf.writestr("[Content_Types].xml", content_types)
            zf.writestr("_rels/.rels", rels)
            zf.writestr("word/document.xml", doc_xml)
        return buffer.getvalue()

    def _load(self, report_id: uuid.UUID, org_id: uuid.UUID, db: Session) -> tuple[ComplianceReport, Organization, list[ComplianceReportSection]]:
        report = ReportRepository(db).get_report(report_id)
        if report is None or report.organization_id != org_id:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Report not found")

        org = db.get(Organization, org_id)
        if org is None:
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="Organization not found")

        sections = ReportRepository(db).list_sections(org_id, report_id)
        return report, org, sections

    @staticmethod
    def _extract_content(report: ComplianceReport, sections: list[ComplianceReportSection]) -> dict[str, Any]:
        if isinstance(report.inputs_summary_json, dict) and report.inputs_summary_json:
            return report.inputs_summary_json

        content = report.content_json if isinstance(report.content_json, dict) else {}
        if content and any(key != "sections" for key in content):
            return content

        assembled: dict[str, Any] = {}
        for section in sections:
            assembled[section.section_key] = section.data_json if section.data_json else section.body_markdown
        return assembled

    @staticmethod
    def _stringify(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, (dict, list)):
            return str(value)
        return str(value)

    def render(self, report_id: uuid.UUID, org_id: uuid.UUID, db: Session) -> bytes:
        report, org, sections = self._load(report_id, org_id, db)
        content = self._extract_content(report, sections)

        report_date = (report.generated_at if report.generated_at else datetime.now(UTC)).date().isoformat()
        title = report.title or report.report_type.replace("_", " ").title()

        try:
            from docx import Document  # type: ignore
        except Exception:  # noqa: BLE001
            return self._fallback_docx_bytes(f"CompliVibe report {report_id}")

        doc = Document()
        doc.add_heading(org.name, 0)
        doc.add_heading(title, 1)
        doc.add_paragraph(f"Generated: {report_date}")
        doc.add_page_break()

        for key, value in content.items():
            if key == "caveat":
                continue
            doc.add_heading(ReportSectionMapper.get_title(key), 2)
            if isinstance(value, dict):
                for k, v in value.items():
                    doc.add_paragraph(f"{str(k).replace('_', ' ').title()}: {self._stringify(v)}")
            elif isinstance(value, list):
                for item in value:
                    doc.add_paragraph(self._stringify(item), style="List Bullet")
            elif isinstance(value, (int, float, str)):
                doc.add_paragraph(self._stringify(value))
            else:
                doc.add_paragraph(self._stringify(value))

        caveat_text = self._stringify(content.get("caveat")) or REPORT_CAVEAT
        doc.add_heading("Disclaimer", 2)
        doc.add_paragraph(caveat_text)

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()
