"""Output-correctness coverage for the export renderers.

The pre-existing test (tests/unit/test_pdf_word_export_sprint1_p2.py) only asserts
the HTTP export endpoints return 200 with a %PDF / PK magic-number and the right
content-type -- it never opens the produced documents to confirm the expected
CONTENT actually made it into the file body.

This module calls the renderer services directly at the unit level and then
re-opens the produced bytes:
  * DOCX  -> python-docx `Document(BytesIO(...))`, asserting headings / paragraph
             text / bullet items are present in the document body.
  * PDF   -> `%PDF` magic + pypdf text extraction, asserting expected text is on
             the page (a real text layer, not just a valid container).
  * ExportContentBuilder -> assert it assembles the expected sections/rows from a
             real seeded entity into the structured content the renderers consume.
"""

from __future__ import annotations

import io
import uuid
from datetime import UTC, datetime

from docx import Document as DocxDocument
from pypdf import PdfReader

from app.compliance.renderers.docx_report_renderer import DocxReportRenderer
from app.compliance.renderers.pdf_report_renderer import PDFReportRenderer
from app.exports.services.docx_renderer import DocxRenderer
from app.exports.services.export_content_builder import (
    ExportBranding,
    ExportContentBuilder,
    ExportDocumentContent,
    ExportSection,
)
from app.exports.services.pdf_renderer import PDFRenderer
from app.models.compliance_report import ComplianceReport
from app.models.control import Control
from app.models.organization import Organization
from tests.helpers.auth_org import bootstrap_org_user


# --------------------------------------------------------------------------- #
# Deterministic payload helpers (no DB required)
# --------------------------------------------------------------------------- #
def _sample_content() -> ExportDocumentContent:
    return ExportDocumentContent(
        organization_id=uuid.uuid4(),
        title="Control Export: Encryption At Rest",
        subtitle="Unit-test subtitle line",
        generated_at=datetime(2026, 7, 17, 12, 0, tzinfo=UTC),
        entity_type="control",
        entity_id=uuid.uuid4(),
        report_type=None,
        branding=ExportBranding(
            organization_id=uuid.uuid4(),
            company_display_name="Acme Compliance Ltd",
            footer_text="Confidential - Acme Internal",
            primary_color_hex="#AA1133",
        ),
        sections=[
            ExportSection(
                title="Control",
                rows=[
                    ("Title", "Encryption At Rest"),
                    ("Criticality", "critical"),
                    ("Status", "planned"),
                ],
            ),
            ExportSection(
                title="Insights",
                items=[
                    "Criticality is critical but status is 'planned'.",
                    "No testing procedure is documented.",
                ],
                paragraphs=["Reviewer note: escalate before audit."],
            ),
        ],
    )


def _pdf_text(data: bytes) -> str:
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


# --------------------------------------------------------------------------- #
# DOCX export renderer -> open with python-docx and read the body
# --------------------------------------------------------------------------- #
def test_docx_export_renderer_body_contains_expected_content():
    data = DocxRenderer().render(_sample_content())
    assert data[:2] == b"PK"  # valid zip/docx container

    doc = DocxDocument(io.BytesIO(data))
    all_text = [p.text for p in doc.paragraphs]
    joined = "\n".join(all_text)

    # Headings the renderer writes (company name @0, title @1, section titles @2).
    heading_texts = {p.text for p in doc.paragraphs if p.style.name.startswith("Heading") or p.style.name == "Title"}
    assert "Acme Compliance Ltd" in heading_texts
    assert "Control Export: Encryption At Rest" in heading_texts
    assert "Control" in heading_texts
    assert "Insights" in heading_texts

    # Subtitle + generated timestamp paragraphs.
    assert "Unit-test subtitle line" in joined
    assert "Generated at: 2026-07-17T12:00:00+00:00" in joined

    # Section rows rendered as "key: value" paragraphs.
    assert "Title: Encryption At Rest" in joined
    assert "Criticality: critical" in joined
    assert "Status: planned" in joined

    # Section free-text paragraph.
    assert "Reviewer note: escalate before audit." in joined

    # Bullet insight items rendered with the List Bullet style.
    bullet_texts = {p.text for p in doc.paragraphs if p.style.name == "List Bullet"}
    assert "Criticality is critical but status is 'planned'." in bullet_texts
    assert "No testing procedure is documented." in bullet_texts

    # Footer text present.
    assert "Confidential - Acme Internal" in joined


# --------------------------------------------------------------------------- #
# PDF export renderer -> valid %PDF + extractable text layer
# --------------------------------------------------------------------------- #
def test_pdf_export_renderer_reportlab_has_valid_text_layer():
    # Exercise the reportlab path directly so the text layer is deterministic and
    # independent of the weasyprint HTML->PDF backend.
    data = PDFRenderer._render_with_reportlab(_sample_content())
    assert data[:4] == b"%PDF"
    assert len(data) > 1000

    text = _pdf_text(data)
    assert "Control Export: Encryption At Rest" in text
    assert "Acme Compliance Ltd" in text
    # Table header + row cell values from the section rows.
    assert "Field" in text and "Value" in text
    assert "Encryption At Rest" in text
    assert "critical" in text
    # Bullet insight text and footer.
    assert "No testing procedure is documented." in text
    assert "Confidential - Acme Internal" in text


def test_pdf_export_renderer_default_render_produces_valid_pdf():
    # The public render() prefers weasyprint and falls back to reportlab; either
    # way it must be a valid, non-trivial PDF for a representative payload.
    data = PDFRenderer().render(_sample_content())
    assert data[:4] == b"%PDF"
    assert len(data) > 1000
    # Both backends embed a real text layer; the title must be extractable.
    text = _pdf_text(data)
    assert "Encryption At Rest" in text


# --------------------------------------------------------------------------- #
# ExportContentBuilder -> assembles expected sections/rows from a real entity
# --------------------------------------------------------------------------- #
def test_export_content_builder_assembles_control_sections(client, db_session):
    owner = bootstrap_org_user(client, email_prefix="renderer-builder")
    org_id = uuid.UUID(owner["organization_id"])
    user_id = uuid.UUID(owner["user_id"])

    control = Control(
        organization_id=org_id,
        title="Access Review Control",
        description="Quarterly access recertification",
        control_type="preventive",
        status="planned",
        criticality="critical",
        created_by_user_id=user_id,
        owner_user_id=user_id,
    )
    db_session.add(control)
    db_session.flush()

    content = ExportContentBuilder(db_session).build_control(org_id, control.id)

    assert isinstance(content, ExportDocumentContent)
    assert content.entity_type == "control"
    assert content.entity_id == control.id
    assert content.title == "Control Export: Access Review Control"

    # The first section is the field table; assert the row key/value tuples.
    control_section = content.sections[0]
    assert control_section.title == "Control"
    rows = dict(control_section.rows)
    assert rows["Title"] == "Access Review Control"
    assert rows["Description"] == "Quarterly access recertification"
    assert rows["Control Type"] == "preventive"
    assert rows["Status"] == "planned"
    assert rows["Criticality"] == "critical"

    # critical + not-effective should have produced an Insights section + flag.
    assert "critical_control_not_effective" in content.context_flags
    insight_titles = [s.title for s in content.sections]
    assert "Insights" in insight_titles
    insights = next(s for s in content.sections if s.title == "Insights").items
    assert any("not yet operating effectively" in i for i in insights)

    # And the assembled content must round-trip through the renderers with the
    # entity data intact.
    docx_bytes = DocxRenderer().render(content)
    doc = DocxDocument(io.BytesIO(docx_bytes))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Title: Access Review Control" in joined

    pdf_bytes = PDFRenderer._render_with_reportlab(content)
    assert pdf_bytes[:4] == b"%PDF"
    assert "Access Review Control" in _pdf_text(pdf_bytes)


# --------------------------------------------------------------------------- #
# Compliance report renderers (app/compliance/renderers) -> full render() path
# --------------------------------------------------------------------------- #
def _seed_report(db_session, org_id, user_id) -> ComplianceReport:
    report = ComplianceReport(
        organization_id=org_id,
        report_type="quarterly_posture",
        title="Q3 Compliance Posture",
        status="generated",
        generated_at=datetime(2026, 7, 17, tzinfo=UTC),
        generated_by_user_id=user_id,
        content_json={
            "score": {"overall": 82, "trend": "up"},
            "risks_summary": ["Vendor onboarding gap", "MFA rollout incomplete"],
            "narrative": "Posture improved this quarter driven by control uplift.",
            "caveat": "Custom disclaimer for this report.",
        },
    )
    db_session.add(report)
    db_session.flush()
    return report


def test_compliance_docx_report_renderer_body_contains_expected_content(client, db_session):
    owner = bootstrap_org_user(client, email_prefix="renderer-docx-report")
    org_id = uuid.UUID(owner["organization_id"])
    user_id = uuid.UUID(owner["user_id"])
    report = _seed_report(db_session, org_id, user_id)

    data = DocxReportRenderer().render(report.id, org_id, db_session)
    assert data[:2] == b"PK"

    doc = DocxDocument(io.BytesIO(data))
    joined = "\n".join(p.text for p in doc.paragraphs)

    # Org name (heading 0) + report title (heading 1) must both appear.
    org = db_session.get(Organization, org_id)
    assert org.name in joined
    assert "Q3 Compliance Posture" in joined
    # Section keys mapped to human titles by ReportSectionMapper.
    assert "Compliance Score" in joined  # score -> "Compliance Score"
    assert "Top Open Risks" in joined  # risks_summary -> "Top Open Risks"
    assert "Executive Summary" in joined  # narrative -> "Executive Summary"
    # Dict field rendered "Key: Value".
    assert "Overall: 82" in joined
    # List rendered as bullets.
    bullet_texts = {p.text for p in doc.paragraphs if p.style.name == "List Bullet"}
    assert "Vendor onboarding gap" in bullet_texts
    assert "MFA rollout incomplete" in bullet_texts
    # Narrative scalar text.
    assert "Posture improved this quarter driven by control uplift." in joined
    # Custom caveat under Disclaimer heading.
    assert "Disclaimer" in joined
    assert "Custom disclaimer for this report." in joined


def test_compliance_pdf_report_renderer_has_valid_text_layer(client, db_session):
    owner = bootstrap_org_user(client, email_prefix="renderer-pdf-report")
    org_id = uuid.UUID(owner["organization_id"])
    user_id = uuid.UUID(owner["user_id"])
    report = _seed_report(db_session, org_id, user_id)

    data = PDFReportRenderer().render(report.id, org_id, db_session)
    assert data[:4] == b"%PDF"
    assert len(data) > 1000

    text = _pdf_text(data)
    assert "Q3 Compliance Posture" in text
    assert "Compliance Score" in text
    assert "Top Open Risks" in text
    assert "Vendor onboarding gap" in text
    assert "Posture improved this quarter driven by control uplift." in text
    assert "Custom disclaimer for this report." in text


def test_compliance_pdf_coverage_matrix_has_valid_text_layer():
    # render_coverage_matrix takes a plain payload -- no DB needed.
    payload = {
        "total_obligations": 3,
        "covered": 1,
        "partial": 1,
        "uncovered": 1,
        "coverage_pct": 33,
        "sections": [
            {
                "section_title": "Access Control",
                "obligations": [
                    {
                        "reference": "AC-1",
                        "title": "Account Management Policy",
                        "controls_count": 2,
                        "evidence_count": 3,
                        "coverage_status": "covered",
                    },
                    {
                        "reference": "AC-2",
                        "title": "Least Privilege Enforcement",
                        "controls_count": 1,
                        "evidence_count": 0,
                        "coverage_status": "partial",
                    },
                ],
            }
        ],
    }
    data = PDFReportRenderer().render_coverage_matrix(
        org_name="Acme Compliance Ltd",
        framework_name="NIST 800-53",
        matrix_payload=payload,
    )
    assert data[:4] == b"%PDF"
    text = _pdf_text(data)
    assert "Framework Coverage Matrix" in text
    assert "NIST 800-53" in text
    assert "Access Control" in text
    assert "Account Management Policy" in text
    assert "Least Privilege Enforcement" in text
    # Summary line numbers.
    assert "Coverage: 33%" in text


# --------------------------------------------------------------------------- #
# Fallback renderers (used when python-docx / reportlab unavailable at runtime)
# --------------------------------------------------------------------------- #
def test_docx_report_fallback_is_openable_and_carries_text():
    data = DocxReportRenderer._fallback_docx_bytes("Fallback report body text")
    assert data[:2] == b"PK"
    doc = DocxDocument(io.BytesIO(data))
    joined = "\n".join(p.text for p in doc.paragraphs)
    assert "Fallback report body text" in joined


def test_pdf_report_fallback_is_valid_pdf():
    data = PDFReportRenderer._fallback_pdf_bytes("Fallback report body text")
    assert data[:4] == b"%PDF"
    # The fallback is a hand-built minimal PDF; pypdf may not extract the text
    # layer reliably, so assert container validity + that the literal string is
    # embedded in the raw stream.
    assert b"Fallback report body text" in data
    reader = PdfReader(io.BytesIO(data))
    assert len(reader.pages) == 1
